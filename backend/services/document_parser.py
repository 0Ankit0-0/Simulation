import os
import logging
from typing import Optional, Dict, Any
import easyocr
import cv2
import numpy as np
from PIL import Image
import pytesseract
from pdf2image import convert_from_path
import docx
import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

class DocumentParser:
    """Enhanced document parser with OCR capabilities"""
    
    def __init__(self):
        """Initialize the parser with OCR models"""
        try:
            # Initialize EasyOCR reader (loads once, ~500MB)
            # Use GPU if available, fallback to CPU
            self.ocr_reader = easyocr.Reader(['en'], gpu=True)
            logger.info("EasyOCR initialized successfully")
        except Exception as e:
            logger.warning(f"EasyOCR initialization failed: {e}")
            self.ocr_reader = None
        
        # Fallback to Tesseract if EasyOCR fails
        try:
            # Test Tesseract availability
            pytesseract.get_tesseract_version()
            self.tesseract_available = True
            logger.info("Tesseract available as fallback")
        except Exception as e:
            logger.warning(f"Tesseract not available: {e}")
            self.tesseract_available = False

    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from various file formats
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.pdf':
                return self._extract_from_pdf(file_path)
            elif file_ext == '.docx':
                return self._extract_from_docx(file_path)
            elif file_ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
                return self._extract_from_image(file_path)
            elif file_ext == '.txt':
                return self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return {
                'text': '',
                'confidence': 0.0,
                'method': 'error',
                'error': str(e)
            }

    def _extract_from_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF using PyMuPDF + OCR for images"""
        text_content = []
        total_confidence = 0.0
        page_count = 0
        
        try:
            # First try to extract text directly
            doc = fitz.open(file_path)
            direct_text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                direct_text += page_text + "\n"
            
            doc.close()
            
            # If direct text extraction yields substantial content, use it
            if len(direct_text.strip()) > 100:
                return {
                    'text': direct_text.strip(),
                    'confidence': 0.95,
                    'method': 'direct_pdf_extraction',
                    'pages': len(doc)
                }
            
            # If not much text found, use OCR on PDF pages
            logger.info("PDF has minimal text, using OCR...")
            return self._ocr_pdf_pages(file_path)
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            return self._ocr_pdf_pages(file_path)

    def _ocr_pdf_pages(self, file_path: str) -> Dict[str, Any]:
        """Convert PDF pages to images and OCR them"""
        try:
            # Convert PDF to images
            images = convert_from_path(file_path, dpi=300, fmt='jpeg')
            
            all_text = []
            total_confidence = 0.0
            
            for i, image in enumerate(images):
                # Convert PIL image to numpy array for processing
                img_array = np.array(image)
                
                # Extract text from image
                result = self._ocr_image_array(img_array)
                all_text.append(result['text'])
                total_confidence += result['confidence']
                
                logger.info(f"Processed PDF page {i+1}/{len(images)}")
            
            avg_confidence = total_confidence / len(images) if images else 0.0
            
            return {
                'text': '\n\n'.join(all_text),
                'confidence': avg_confidence,
                'method': 'pdf_ocr',
                'pages': len(images)
            }
            
        except Exception as e:
            logger.error(f"Error in PDF OCR: {e}")
            return {
                'text': '',
                'confidence': 0.0,
                'method': 'pdf_ocr_failed',
                'error': str(e)
            }

    def _extract_from_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        try:
            doc = docx.Document(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            full_text = '\n'.join(text_content)
            
            return {
                'text': full_text,
                'confidence': 0.98,
                'method': 'docx_direct',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            return {
                'text': '',
                'confidence': 0.0,
                'method': 'docx_failed',
                'error': str(e)
            }

    def _extract_from_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image files using OCR"""
        try:
            # Load and preprocess image
            image = cv2.imread(file_path)
            if image is None:
                raise ValueError("Could not load image")
            
            return self._ocr_image_array(image)
            
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            return {
                'text': '',
                'confidence': 0.0,
                'method': 'image_ocr_failed',
                'error': str(e)
            }

    def _ocr_image_array(self, image_array: np.ndarray) -> Dict[str, Any]:
        """Perform OCR on image array"""
        # Preprocess image for better OCR
        processed_image = self._preprocess_image(image_array)
        
        # Try EasyOCR first
        if self.ocr_reader:
            try:
                results = self.ocr_reader.readtext(processed_image)
                
                # Extract text and calculate average confidence
                text_parts = []
                confidences = []
                
                for (bbox, text, confidence) in results:
                    if confidence > 0.3:  # Filter low confidence results
                        text_parts.append(text)
                        confidences.append(confidence)
                
                full_text = ' '.join(text_parts)
                avg_confidence = sum(confidences) / len(confidences) if confidences else 0.0
                
                return {
                    'text': full_text,
                    'confidence': avg_confidence,
                    'method': 'easyocr',
                    'detected_regions': len(results)
                }
                
            except Exception as e:
                logger.warning(f"EasyOCR failed, trying Tesseract: {e}")
        
        # Fallback to Tesseract
        if self.tesseract_available:
            try:
                # Convert BGR to RGB for Tesseract
                rgb_image = cv2.cvtColor(processed_image, cv2.COLOR_BGR2RGB)
                text = pytesseract.image_to_string(rgb_image)
                
                return {
                    'text': text.strip(),
                    'confidence': 0.7,  # Tesseract doesn't provide confidence easily
                    'method': 'tesseract',
                }
                
            except Exception as e:
                logger.error(f"Tesseract OCR failed: {e}")
        
        # If all OCR methods fail
        return {
            'text': '',
            'confidence': 0.0,
            'method': 'ocr_all_failed',
            'error': 'All OCR methods failed'
        }

    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Preprocess image for better OCR results"""
        # Convert to grayscale
        gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        
        # Apply denoising
        denoised = cv2.fastNlMeansDenoising(gray)
        
        # Apply adaptive thresholding
        thresh = cv2.adaptiveThreshold(
            denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
        )
        
        # Resize if image is too small (improves OCR accuracy)
        height, width = thresh.shape
        if height < 300 or width < 300:
            scale_factor = max(300 / height, 300 / width)
            new_width = int(width * scale_factor)
            new_height = int(height * scale_factor)
            thresh = cv2.resize(thresh, (new_width, new_height), interpolation=cv2.INTER_CUBIC)
        
        return cv2.cvtColor(thresh, cv2.COLOR_GRAY2BGR)

    def _extract_from_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            return {
                'text': text,
                'confidence': 1.0,
                'method': 'direct_text',
                'char_count': len(text)
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                
                return {
                    'text': text,
                    'confidence': 0.9,
                    'method': 'direct_text_latin1',
                    'char_count': len(text)
                }
            except Exception as e:
                logger.error(f"Error reading text file: {e}")
                return {
                    'text': '',
                    'confidence': 0.0,
                    'method': 'text_failed',
                    'error': str(e)
                }

# Global parser instance
document_parser = DocumentParser()

def extract_text_from_file(file_path: str) -> Dict[str, Any]:
    """
    Main function to extract text from any supported file
    This is the function your existing code calls
    """
    return document_parser.extract_text_from_file(file_path)